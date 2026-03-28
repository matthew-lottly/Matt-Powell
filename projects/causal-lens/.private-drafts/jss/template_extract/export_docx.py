from __future__ import annotations

import re
from pathlib import Path
from typing import cast

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle

ROOT = Path(__file__).resolve().parent
TEX_PATH = ROOT / "causal-lens.tex"
BBL_PATH = ROOT / "causal-lens.bbl"
PROJECT_ROOT = ROOT.parents[2]
DOCS_DIR = PROJECT_ROOT / "docs"
DOCX_PATH = DOCS_DIR / "causal-lens-review-copy.docx"

DOCX_TEXT_REPLACEMENTS = {
    "replications/run_all.py": "the reproducibility runner",
    "replications/outputs/": "the replication outputs directory",
    "replications/outputs": "the replication outputs directory",
    "outputs/paper/": "the paper outputs directory",
    "outputs/paper": "the paper outputs directory",
    "Placebo/falsification": "Placebo and falsification",
    "Double/Debiased": "Double or debiased",
    "Double/debiased": "Double or debiased",
}


def load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def simplify_math(text: str) -> str:
    replacements = {
        r"\tau": "\u03c4",
        r"\Gamma": "\u0393",
        r"\geq": "\u2265",
        r"\leq": "\u2264",
        r"\neq": "\u2260",
        r"\times": "\u00d7",
        r"\pm": "\u00b1",
        r"\top": "\u1d40",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"\\hat\{([^{}]+)\}", lambda m: f"{m.group(1)}\u0302", text)
    text = re.sub(r"\\mathrm\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\text\{([^{}]*)\}", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def simplify_latex(text: str) -> str:
    text = text.replace(r"\$", "__JSS_DOLLAR__")
    text = re.sub(r"\\+%", "%", text)

    replacements = {
        r"\%": "%",
        r"\_": "_",
        r"\&": "&",
        r"\\": " ",
        r"~": " ",
        "``": '"',
        "''": '"',
        "---": "--",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"\\['\^\"`~=\.]\{?([A-Za-z])\}?", r"\1", text)
    text = re.sub(r"\\[Hcdbuvtk]\{([A-Za-z])\}", r"\1", text)
    text = text.replace(r"\L", "L")
    text = re.sub(r"\\(proglang|pkg|code|emph|textbf|texttt|mathrm)\{([^{}]*)\}", r"\2", text)
    text = re.sub(r"\$([^$]+)\$", lambda m: simplify_math(m.group(1)), text)
    text = re.sub(r"\\url\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\doi\{([^{}]+)\}", r"doi: \1", text)
    text = re.sub(r"\\enquote\{", "", text)
    text = re.sub(r"\\enquote\{([^{}]+)\}", r'"\1"', text)
    text = re.sub(r"\\newblock", " ", text)
    text = re.sub(r"\\natexlab\{[^}]+\}", "", text)
    text = re.sub(r"\\label\{[^}]+\}", "", text)
    text = re.sub(r"\\centering", "", text)
    text = re.sub(r"\\footnotesize", "", text)
    text = re.sub(r"\\setlength\{[^}]+\}\{[^}]+\}", "", text)
    text = re.sub(r"\\urlprefix", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"(?<=\w)-\s+(?=\w)", "-", text)
    text = text.replace("__JSS_DOLLAR__", "$")
    text = re.sub(r"(?<=\d)and(?=\d)", " and ", text)
    text = re.sub(r"\bp value\b", "p-value", text)
    text = re.sub(r"\bF statistic\b", "F-statistic", text)
    text = re.sub(r"\s+", " ", text)
    for old, new in DOCX_TEXT_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text.strip()


def parse_citation_map(bbl_text: str) -> tuple[dict[str, str], dict[str, str], list[str]]:
    citep_map: dict[str, str] = {}
    citet_map: dict[str, str] = {}
    refs: list[str] = []

    pattern = re.compile(
        r"\\bibitem\[(?P<display>.*?)\]\{(?P<key>[^}]+)\}(?P<body>.*?)(?=\\bibitem\[|\\end\{thebibliography\})",
        re.DOTALL,
    )
    for match in pattern.finditer(bbl_text):
        display = match.group("display")
        key = match.group("key")
        body = match.group("body")
        author = display.split("(")[0]
        year_match = re.search(r"\((\d{4})", display)
        year = year_match.group(1) if year_match else "n.d."
        author = simplify_latex(author).replace(" et al.", " et al.")
        author = author.replace(" et~al.", " et al.")
        author = author.strip()
        citep_map[key] = f"{author} {year}"
        citet_map[key] = f"{author} ({year})"

        entry_text = simplify_latex(body)
        if entry_text:
            refs.append(entry_text)
    return citep_map, citet_map, refs


def humanize_citation_key(key: str) -> str:
    key = key.strip()
    match = re.match(r"(?P<authors>.+?)(?P<year>\d{4})$", key)
    if not match:
        return key

    authors = match.group("authors")
    year = match.group("year")
    authors = re.sub(r"EtAl$", " et al.", authors)

    if authors.endswith(" et al."):
        lead = authors[: -len(" et al.")]
        lead = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", lead)
        authors = f"{lead} et al."
    else:
        parts = re.findall(r"[A-Z][a-z']*(?:-[A-Z][a-z']*)?|[A-Z]{2,}(?=[A-Z][a-z]|$)", authors)
        if parts:
            if len(parts) == 2 and parts[0] == "Mc":
                authors = f"Mc{parts[1]}"
            elif len(parts) == 2:
                authors = f"{parts[0]} and {parts[1]}"
            elif len(parts) > 2:
                authors = ", ".join(parts[:-1]) + f", and {parts[-1]}"
            else:
                authors = parts[0]

    return f"{authors} {year}".strip()


def replace_citations(
    text: str,
    citep_map: dict[str, str],
    citet_map: dict[str, str],
    ref_map: dict[str, str],
) -> str:
    def repl_citep(match: re.Match[str]) -> str:
        keys = [k.strip() for k in match.group(1).split(",")]
        return "(" + "; ".join(citep_map.get(k, humanize_citation_key(k)) for k in keys) + ")"

    def repl_citet(match: re.Match[str]) -> str:
        keys = [k.strip() for k in match.group(1).split(",")]
        return "; ".join(citet_map.get(k, humanize_citation_key(k)) for k in keys)

    text = re.sub(r"\\citep\{([^}]+)\}", repl_citep, text)
    text = re.sub(r"\\citet\{([^}]+)\}", repl_citet, text)
    text = re.sub(r"\\cite\{([^}]+)\}", repl_citet, text)
    text = re.sub(r"\\ref\{([^}]+)\}", lambda m: ref_map.get(m.group(1), m.group(1)), text)
    return text


def clean_text(
    text: str,
    citep_map: dict[str, str],
    citet_map: dict[str, str],
    ref_map: dict[str, str],
) -> str:
    text = replace_citations(text, citep_map, citet_map, ref_map)
    return simplify_latex(text)


def extract_braced_block(text: str, command: str) -> str:
    start = text.find(command)
    if start == -1:
        return ""
    start = text.find("{", start)
    depth = 0
    for index in range(start, len(text)):
        if text[index] == "{":
            depth += 1
        elif text[index] == "}":
            depth -= 1
            if depth == 0:
                return text[start + 1:index]
    return ""


def extract_label(text: str) -> str | None:
    match = re.search(r"\\label\{([^}]+)\}", text)
    return match.group(1) if match else None


def build_reference_map(tex_text: str) -> dict[str, str]:
    ref_map: dict[str, str] = {}
    figure_count = 0
    table_count = 0
    lines = tex_text.splitlines()
    in_figure = False
    in_table = False
    block_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.strip()

        if in_figure:
            block_lines.append(raw_line)
            if line.startswith(r"\end{figure}"):
                figure_count += 1
                label = extract_label("\n".join(block_lines))
                if label:
                    ref_map[label] = str(figure_count)
                block_lines = []
                in_figure = False
            continue

        if in_table:
            block_lines.append(raw_line)
            if line.startswith(r"\end{table}"):
                table_count += 1
                label = extract_label("\n".join(block_lines))
                if label:
                    ref_map[label] = str(table_count)
                block_lines = []
                in_table = False
            continue

        if line.startswith(r"\begin{figure}"):
            in_figure = True
            block_lines = [raw_line]
            continue

        if line.startswith(r"\begin{table}"):
            in_table = True
            block_lines = [raw_line]
            continue

        if line.startswith(r"\input{"):
            input_path = line[len(r"\input{"):-1]
            table_text = load_text((ROOT / input_path).resolve())
            label = extract_label(table_text)
            table_count += 1
            if label:
                ref_map[label] = str(table_count)

    return ref_map


def format_paragraph(paragraph, *, indent: bool = True, centered: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(0.3) if indent else Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def add_body_paragraph(document: DocxDocument, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph(text)
    format_paragraph(paragraph, indent=indent)


def add_caption(document: DocxDocument, prefix: str, caption_text: str) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, indent=False, centered=False)
    run_prefix = paragraph.add_run(f"{prefix}. ")
    run_prefix.bold = True
    run_caption = paragraph.add_run(caption_text)
    run_caption.italic = True


def parse_table_text(
    table_text: str,
    document: DocxDocument,
    citep_map: dict[str, str],
    citet_map: dict[str, str],
    ref_map: dict[str, str],
) -> None:
    caption_text = extract_braced_block(table_text, r"\caption")
    label = extract_label(table_text)

    rows: list[list[str]] = []
    in_tabular = False
    for raw_line in table_text.splitlines():
        line = raw_line.strip()
        if line.startswith(r"\begin{tabular"):
            in_tabular = True
            continue
        if line.startswith(r"\end{tabular"):
            in_tabular = False
            continue
        if not in_tabular or not line or line == r"\hline":
            continue
        if "&" not in line:
            continue
        line = re.sub(r"\\\\\s*$", "", line).strip()
        cells = [clean_text(cell.strip(), citep_map, citet_map, ref_map) for cell in line.split("&")]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, indent=False)
                if row_idx == 0:
                    for run in paragraph.runs:
                        run.bold = True

    if caption_text:
        caption = clean_text(caption_text, citep_map, citet_map, ref_map)
        number = ref_map.get(label, "") if label else ""
        prefix = f"Table {number}".strip()
        add_caption(document, prefix or "Table", caption)
    document.add_paragraph()


def add_code_block(document: DocxDocument, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = document.add_paragraph()
        format_paragraph(paragraph, indent=False)
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line.rstrip())
        font = run.font
        font.name = "Courier New"
        font.size = Pt(9)
    document.add_paragraph()


def add_figure(
    document: DocxDocument,
    figure_block: str,
    citep_map: dict[str, str],
    citet_map: dict[str, str],
    ref_map: dict[str, str],
) -> None:
    path_match = re.search(r"\\includegraphics\[[^]]*\]\{([^}]+)\}", figure_block)
    caption_text = extract_braced_block(figure_block, r"\caption")
    label = extract_label(figure_block)
    if path_match:
        rel_path = path_match.group(1)
        image_path = (ROOT / rel_path).resolve()
        if image_path.suffix.lower() == ".pdf":
            png_candidate = image_path.with_suffix(".png")
            if png_candidate.exists():
                image_path = png_candidate
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_text:
        caption = clean_text(caption_text, citep_map, citet_map, ref_map)
        number = ref_map.get(label, "") if label else ""
        prefix = f"Figure {number}".strip()
        add_caption(document, prefix or "Figure", caption)
    document.add_paragraph()


def build_docx() -> Path:
    tex_text = load_text(TEX_PATH)
    bbl_text = load_text(BBL_PATH)
    citep_map, citet_map, refs = parse_citation_map(bbl_text)
    ref_map = build_reference_map(tex_text)

    document = create_document()
    normal_style = cast(ParagraphStyle, document.styles["Normal"])
    normal = normal_style.font
    normal.name = "Times New Roman"
    normal.size = Pt(11)
    normal_style.paragraph_format.first_line_indent = Inches(0.3)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
        heading_style = cast(ParagraphStyle, document.styles[style_name])
        heading_style.font.name = "Times New Roman"
        heading_style.font.size = Pt(size)
        heading_style.paragraph_format.first_line_indent = Inches(0)

    title = simplify_latex(extract_braced_block(tex_text, r"\title"))
    author_block = extract_braced_block(tex_text, r"\author")
    author_lines = [simplify_latex(part) for part in re.split(r"\\\\", author_block) if simplify_latex(part)]
    address_block = extract_braced_block(tex_text, r"\Address")
    address_lines = [simplify_latex(part) for part in re.split(r"\\\\", address_block) if simplify_latex(part)]
    address_lines = [line for line in address_lines if line not in author_lines]
    abstract = clean_text(extract_braced_block(tex_text, r"\Abstract"), citep_map, citet_map, ref_map)
    keywords = clean_text(extract_braced_block(tex_text, r"\Keywords"), citep_map, citet_map, ref_map)

    p = document.add_paragraph()
    format_paragraph(p, indent=False, centered=True)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    for author_line in author_lines:
        p = document.add_paragraph()
        format_paragraph(p, indent=False, centered=True)
        p.add_run(author_line)

    for address_line in address_lines:
        p = document.add_paragraph()
        format_paragraph(p, indent=False, centered=True)
        p.add_run(address_line)

    p = document.add_paragraph()
    format_paragraph(p, indent=False, centered=True)
    p.add_run("Personal review copy generated from the JSS LaTeX source.")

    heading = document.add_heading("Abstract", level=1)
    format_paragraph(heading, indent=False)
    add_body_paragraph(document, abstract, indent=False)
    add_body_paragraph(document, f"Keywords: {keywords}", indent=False)

    body = tex_text.split(r"\begin{document}", 1)[1].split(r"\bibliography{causal-lens}", 1)[0]
    lines = body.splitlines()

    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []
    in_figure = False
    figure_lines: list[str] = []
    in_table = False
    table_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        text = " ".join(line.strip() for line in paragraph_buffer if line.strip())
        text = clean_text(text, citep_map, citet_map, ref_map)
        if text:
            add_body_paragraph(document, text)
        paragraph_buffer = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line or line.startswith("%%"):
            if not in_code and not in_figure and not in_table:
                flush_paragraph()
            continue

        if in_code:
            if line.startswith(r"\end{Code}"):
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw_line)
            continue

        if in_table:
            table_lines.append(raw_line)
            if line.startswith(r"\end{table}"):
                parse_table_text("\n".join(table_lines), document, citep_map, citet_map, ref_map)
                table_lines = []
                in_table = False
            continue

        if in_figure:
            figure_lines.append(raw_line)
            if line.startswith(r"\end{figure}"):
                add_figure(document, "\n".join(figure_lines), citep_map, citet_map, ref_map)
                figure_lines = []
                in_figure = False
            continue

        if line.startswith(r"\begin{Code}"):
            flush_paragraph()
            in_code = True
            continue

        if line.startswith(r"\begin{figure}"):
            flush_paragraph()
            in_figure = True
            figure_lines = [raw_line]
            continue

        if line.startswith(r"\begin{table}"):
            flush_paragraph()
            in_table = True
            table_lines = [raw_line]
            continue

        if line.startswith(r"\section"):
            flush_paragraph()
            title_text = clean_text(extract_braced_block(line, r"\section"), citep_map, citet_map, ref_map)
            if not title_text:
                brace_parts = re.findall(r"\{([^{}]+)\}", line)
                title_text = clean_text(brace_parts[-1], citep_map, citet_map, ref_map) if brace_parts else line
            heading = document.add_heading(title_text, level=1)
            format_paragraph(heading, indent=False)
            continue

        if line.startswith(r"\subsection"):
            flush_paragraph()
            title_text = clean_text(extract_braced_block(line, r"\subsection"), citep_map, citet_map, ref_map)
            heading = document.add_heading(title_text, level=2)
            format_paragraph(heading, indent=False)
            continue

        if line.startswith(r"\input{"):
            flush_paragraph()
            input_path = line[len(r"\input{"):-1]
            table_text = load_text((ROOT / input_path).resolve())
            parse_table_text(table_text, document, citep_map, citet_map, ref_map)
            continue

        if line in {r"\end{document}", r"\begin{document}"}:
            continue

        paragraph_buffer.append(raw_line)

    flush_paragraph()

    heading = document.add_heading("References", level=1)
    format_paragraph(heading, indent=False)
    for ref in refs:
        paragraph = document.add_paragraph(ref)
        format_paragraph(paragraph, indent=False)
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.3)

    document.save(str(DOCX_PATH))
    return DOCX_PATH


if __name__ == "__main__":
    output = build_docx()
    print(output)
